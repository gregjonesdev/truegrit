from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.core.management.base import BaseCommand


states = [
    "California", "Texas", "New York", "Florida", "Illinois",
    "Pennsylvania", "Ohio", "Georgia", "North Carolina", "Michigan",
    "New Jersey", "Virginia", "Washington", "Arizona", "Massachusetts",
    "Tennessee", "Indiana", "Missouri", "Maryland", "Wisconsin",
    "Colorado", "Minnesota", "South Carolina", "Alabama", "Louisiana",
    "Kentucky", "Oregon", "Iowa", "Oklahoma", "Connecticut", "Utah"
]

cities = [
    "Los Angeles", "Houston", "New York City", "Miami", "Chicago",
    "Philadelphia", "Columbus", "Atlanta", "Charlotte", "Detroit",
    "San Francisco", "Seattle", "Boston", "Phoenix", "Nashville",
    "Indianapolis", "St. Louis", "Baltimore", "Milwaukee", "Denver",
    "Portland", "Minneapolis", "Tulsa", "Cleveland", "Pittsburgh",
    "Omaha", "Las Vegas", "Virginia Beach", "Atlanta", "Cincinnati", "Kansas City"
]

ip_addresses = [
    "192.168.1.1", "10.0.0.1", "172.16.0.1", "192.168.0.1", "192.168.1.2",
    "10.0.0.2", "172.16.0.2", "192.168.0.2", "192.168.1.3", "10.0.0.3",
    "172.16.0.3", "192.168.0.3", "192.168.1.4", "10.0.0.4", "172.16.0.4",
    "192.168.0.4", "192.168.1.5", "10.0.0.5", "172.16.0.5", "192.168.0.5",
    "192.168.1.6", "10.0.0.6", "172.16.0.6", "192.168.0.6", "192.168.1.7",
    "10.0.0.7", "172.16.0.7", "192.168.0.7", "192.168.1.8", "10.0.0.8",
    "172.16.0.8"
]

class Command(BaseCommand):  

    def handle(self, *args, **options):

        document = Document()

        # Create a table for the labels
        table = document.add_table(rows=10, cols=3)
        
        # Set the column width
        for cell in table.columns:
            for c in cell.cells:
                c.width = Inches(2)

        # Populate the table with labels
        for i in range(30):
            state = states[i % len(states)]
            city = cities[i % len(cities)]
            ip = ip_addresses[i % len(ip_addresses)]

            # Insert the label text into the corresponding table cell
            cell = table.cell(i // 3, i % 3)
            # cell.text = f"{state}\n{city}\n{ip}"

            p1 = cell.add_paragraph()
            p1.add_run(state).font.size = Pt(12)  # State on the first line

            p2 = cell.add_paragraph()
            p2.add_run(city).font.size = Pt(8)    # City on the second line

            p3 = cell.add_paragraph()
            p3.add_run(ip).font.size = Pt(12) 

            # Adjust font size
            for paragraph in cell.paragraphs:  
                for run in paragraph.runs:
                    run.font.name = 'Arial'  

        # Save the document
        document.save('new_labels.docx')    
